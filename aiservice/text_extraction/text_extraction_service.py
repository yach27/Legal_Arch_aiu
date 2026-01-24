from flask import Flask, request, jsonify
from flask_cors import CORS
import logging
import os

# CRITICAL FIX for PaddleOCR 0xC0000005 Access Violation
# Must be set BEFORE importing numpy/paddle
# os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
# os.environ["OMP_NUM_THREADS"] = "1" 

# CONFIGURATION
USE_PADDLEOCR = True  # Enabled with optimized mobile models for handwriting recognition

import traceback
from datetime import datetime
import PyPDF2
import docx
import zipfile
import xml.etree.ElementTree as ET

# Modern OCR imports
try:
    from pdf2image import convert_from_path
    from PIL import Image, ImageFilter, ImageEnhance
    import numpy as np

    # OpenCV import for advanced preprocessing
    try:
        import cv2
        OPENCV_AVAILABLE = True
    except ImportError:
        OPENCV_AVAILABLE = False

    # PaddleOCR import (primary OCR engine)
    try:
        if USE_PADDLEOCR:
            from paddleocr import PaddleOCR
            PADDLEOCR_AVAILABLE = True
            logger = logging.getLogger(__name__)
            logger.info("PaddleOCR library loaded successfully")
        else:
             PADDLEOCR_AVAILABLE = False
             logger = logging.getLogger(__name__)
             logger.info("PaddleOCR disabled by configuration (reverting to Tesseract)")
    except ImportError as e:
        PADDLEOCR_AVAILABLE = False
        logger = logging.getLogger(__name__)
        logger.warning(f"PaddleOCR not available: {e}")
        # Print full traceback to debug log
        import traceback
        logger.error(traceback.format_exc())

    # Fallback OCR imports
    try:
        import pytesseract
        TESSERACT_AVAILABLE = True

        # Set Tesseract path immediately when imported
        tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info(f"Tesseract configured at: {tesseract_path}")

    except ImportError:
        TESSERACT_AVAILABLE = False
        logger.warning("Tesseract not available as fallback")

except ImportError as e:
    PADDLEOCR_AVAILABLE = False
    TESSERACT_AVAILABLE = False
    OPENCV_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.error(f"OCR libraries not available: {e}")
    logger.error("Text extraction will be severely limited")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# STARTUP VERIFICATION
print("\n" + "="*50)
print("### TEXT EXTRACTION SERVICE RELOADED ###")
print(f"### WORKER COUNT: 8 ###")
print("### PADDLEOCR STATUS: Checking... ###")
print("="*50 + "\n")

app = Flask(__name__)
CORS(app)

class TextExtractor:
    def __init__(self):
        self.paddle_ocr = None

        if PADDLEOCR_AVAILABLE:
            try:
                logger.info("Initializing PaddleOCR model...")

                # Initialize PaddleOCR with lightweight mobile models for speed
                # Mobile models are 2-3x faster than standard models with minimal accuracy loss
                # Models auto-download from HuggingFace to ~/.paddlex/official_models/
                self.paddle_ocr = PaddleOCR(
                    lang='en',
                    text_detection_model_dir='en_PP-OCRv4_mobile_det',  # Lightweight detection model
                    text_recognition_model_dir='en_PP-OCRv4_mobile_rec',  # Lightweight recognition model
                    use_angle_cls=False,
                    enable_mkldnn=True   # Enable MKLDNN for faster CPU inference
                )

                logger.info("PaddleOCR initialized successfully (models auto-downloaded from HuggingFace)")
            except Exception as e:
                logger.error(f"Failed to initialize PaddleOCR: {e}")
                self.paddle_ocr = None
    
    def preprocess_image_for_ocr(self, image):
        """Preprocess image to improve OCR accuracy"""
        try:
            if OPENCV_AVAILABLE:
                # Use OpenCV for advanced preprocessing
                if isinstance(image, Image.Image):
                    image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # Convert to grayscale
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                
                # Apply denoising
                denoised = cv2.medianBlur(gray, 3)
                
                # Apply adaptive thresholding
                thresh = cv2.adaptiveThreshold(
                    denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
                )
                
                # Morphological operations to clean up the image
                kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
                cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
                
                # Convert back to PIL Image
                return Image.fromarray(cleaned)
            else:
                # Use PIL for enhanced preprocessing
                if not isinstance(image, Image.Image):
                    return image

                # Convert to grayscale
                gray = image.convert('L')

                # Enhance sharpness first
                sharpness_enhancer = ImageEnhance.Sharpness(gray)
                sharp = sharpness_enhancer.enhance(1.5)

                # Enhance contrast
                contrast_enhancer = ImageEnhance.Contrast(sharp)
                enhanced = contrast_enhancer.enhance(1.8)

                # Apply adaptive-like threshold using point operations
                # Get image statistics for better threshold
                import statistics
                pixels = list(enhanced.getdata())
                mean_val = statistics.mean(pixels)

                # Use mean-based threshold
                threshold = int(mean_val * 0.9)  # Slightly below mean for better text separation
                processed = enhanced.point(lambda p: 255 if p > threshold else 0)

                return processed
                
        except Exception as e:
            logger.error(f"Image preprocessing failed: {str(e)}")
            return image

    def is_handwritten(self, image):
        """Detect if image likely contains handwriting using edge density and variance"""
        try:
            # Convert to grayscale if needed
            if isinstance(image, Image.Image):
                if image.mode != 'L':
                    gray = image.convert('L')
                else:
                    gray = image
                img_array = np.array(gray)
            else:
                img_array = image
            
            # Resize to standard size for consistent analysis
            if img_array.shape[0] > 500 or img_array.shape[1] > 500:
                scale = 500 / max(img_array.shape)
                new_size = (int(img_array.shape[1] * scale), int(img_array.shape[0] * scale))
                gray_pil = Image.fromarray(img_array)
                gray_pil = gray_pil.resize(new_size, Image.Resampling.LANCZOS)
                img_array = np.array(gray_pil)
            
            # Calculate edge density (handwriting has more irregular edges)
            if OPENCV_AVAILABLE:
                edges = cv2.Canny(img_array, 50, 150)
                edge_density = np.sum(edges > 0) / edges.size
            else:
                # Simple edge detection using PIL
                from PIL import ImageFilter
                gray_pil = Image.fromarray(img_array)
                edges = gray_pil.filter(ImageFilter.FIND_EDGES)
                edge_array = np.array(edges)
                edge_density = np.sum(edge_array > 30) / edge_array.size
            
            # Calculate pixel variance (handwriting has higher variance)
            pixel_variance = np.var(img_array)
            
            # Heuristic thresholds (tuned for legal documents)
            # Handwriting: edge_density > 0.05, variance > 1000
            # Printed text: edge_density < 0.05, variance < 1000
            is_handwritten = edge_density > 0.05 or pixel_variance > 1000
            
            logger.info(f"Handwriting detection: edge_density={edge_density:.4f}, variance={pixel_variance:.1f}, is_handwritten={is_handwritten}")
            return is_handwritten
            
        except Exception as e:
            logger.error(f"Handwriting detection failed: {str(e)}")
            # Default to handwritten (use PaddleOCR) if detection fails
            return True

    def extract_text_with_paddleocr(self, image):
        """Extract text using PaddleOCR (fast and accurate for stamps, handwriting, IDs)"""
        if not self.paddle_ocr:
            return ""

        try:
            logger.info("Using PaddleOCR for text extraction")

            # Ensure image is PIL Image
            if not isinstance(image, Image.Image):
                image = Image.fromarray(image)

            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Convert PIL Image to numpy array for PaddleOCR
            # Resize if too large to prevent crashes and speed up inference
            # Optimized for speed: 1500px is sweet spot for handwriting recognition
            max_dimension = 1500  # Reduced from 2000 for 30% speed improvement
            if max(image.size) > max_dimension:
                scale = max_dimension / max(image.size)
                new_size = (int(image.size[0] * scale), int(image.size[1] * scale))
                logger.info(f"Resizing image from {image.size} to {new_size} for speed")
                image = image.resize(new_size, Image.Resampling.LANCZOS)

            img_array = np.array(image)

            # Run PaddleOCR (cls parameter removed in v3.x)
            # Run PaddleOCR (predict is the new method, ocr is deprecated)
            result = self.paddle_ocr.predict(img_array)

            # Extract text from results
            # PaddleOCR v3 format: result[0] = list of [bbox, (text, confidence)] or None if no text
            extracted_text = []
            if result and result[0]:
                for line in result[0]:
                    if line and len(line) > 1:
                        # line = [bbox, (text, confidence)]
                        text_info = line[1]
                        if isinstance(text_info, (list, tuple)) and len(text_info) >= 2:
                            text = text_info[0]
                            confidence = text_info[1]
                            # Very low threshold to catch faint handwriting, stamps, and IDs
                            if confidence > 0.2:  # 20% - catches even faint blue ink
                                extracted_text.append(text)
                        elif isinstance(text_info, str):
                            # Sometimes just returns string
                            extracted_text.append(text_info)

            if extracted_text:
                full_text = ' '.join(extracted_text)
                logger.info(f"PaddleOCR extracted {len(extracted_text)} text regions")
                return full_text.strip()

            logger.warning("PaddleOCR found no text in image")
            return ""

        except Exception as e:
            logger.error(f"PaddleOCR extraction failed: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            return ""

    def extract_text_with_ocr(self, file_path):
        """Extract text from image-based PDF using PaddleOCR and fallback to Tesseract"""
        if not PADDLEOCR_AVAILABLE and not TESSERACT_AVAILABLE:
            logger.warning("OCR libraries not available, cannot extract from image-based PDF")
            return ""

        try:
            logger.info(f"Attempting OCR extraction from: {file_path}")

            # Convert PDF to images
            # Set poppler path for pdf2image
            poppler_path = r'C:\poppler\poppler-23.08.0\Library\bin'
            if not os.path.exists(poppler_path):
                # Try alternative paths
                alt_paths = [
                    r'C:\poppler\Library\bin',
                    r'C:\poppler\bin',
                    r'C:\Program Files\poppler\bin'
                ]
                for path in alt_paths:
                    if os.path.exists(path):
                        poppler_path = path
                        break
                else:
                    poppler_path = None

            # Increase PIL's image size limit for OCR processing
            from PIL import Image
            Image.MAX_IMAGE_PIXELS = None  # Remove the limit

            # Optimized DPI for handwriting recognition
            # 150 DPI: Perfect balance for handwritten text (2x faster than 300 DPI)
            # Higher DPI doesn't improve handwriting accuracy significantly
            dpi_setting = 150

            # Process fewer pages to speed up load time
            max_pages = 10  # Process up to 10 pages
            pages_to_process = min(total_pages, max_pages) if 'total_pages' in locals() else max_pages
            
            if poppler_path:
                logger.info(f"Using Poppler path: {poppler_path}")
                # Use thread_count=8 to speed up PDF->Image conversion
                images = convert_from_path(file_path, dpi=dpi_setting, fmt='PNG', poppler_path=poppler_path, first_page=1, last_page=max_pages, thread_count=8)
            else:
                logger.warning("Poppler not found, trying without explicit path")
                images = convert_from_path(file_path, dpi=dpi_setting, fmt='PNG', first_page=1, last_page=max_pages, thread_count=8)
            
            # Recalculate based on what we actually got
            pages_to_process = len(images)

            extracted_text_dict = {}
            
            import concurrent.futures
            
            # Function to process a single page (closure to capture self and logging)
            def process_page(page_idx, page_img):
                try:
                    logger.info(f"Processing page {page_idx+1}/{pages_to_process} on thread")
                    p_text = ""
                    
                    # Smart detection: Use Tesseract for printed, PaddleOCR for handwritten
                    if self.paddle_ocr and TESSERACT_AVAILABLE:
                        # Detect if page has handwriting
                        has_handwriting = self.is_handwritten(page_img)
                        
                        if has_handwriting:
                            logger.info(f"Page {page_idx+1}: Detected handwriting, using PaddleOCR")
                            p_text = self.extract_text_with_paddleocr(page_img)
                        else:
                            logger.info(f"Page {page_idx+1}: Detected printed text, using Tesseract (faster)")
                            # Use Tesseract for printed text
                            if OPENCV_AVAILABLE:
                                processed_img = self.preprocess_image_for_ocr(page_img)
                            else:
                                processed_img = page_img.convert('L')
                            custom_config = r'--oem 3 --psm 6'
                            p_text = pytesseract.image_to_string(processed_img, config=custom_config, lang='eng')
                    
                    # Fallback: PaddleOCR only (if Tesseract not available)
                    elif self.paddle_ocr:
                        logger.info(f"Page {page_idx+1}: Using PaddleOCR (Tesseract not available)")
                        p_text = self.extract_text_with_paddleocr(page_img)
                    
                    # Fallback: Tesseract only (if PaddleOCR not available)
                    elif TESSERACT_AVAILABLE:
                        logger.info(f"Page {page_idx+1}: Using Tesseract (PaddleOCR not available)")
                        if OPENCV_AVAILABLE:
                            processed_img = self.preprocess_image_for_ocr(page_img)
                        else:
                            processed_img = page_img.convert('L')
                        custom_config = r'--oem 3 --psm 6'
                        p_text = pytesseract.image_to_string(processed_img, config=custom_config, lang='eng')
                    
                    return page_idx, p_text
                except Exception as ex:
                    logger.error(f"Error on page {page_idx+1}: {ex}")
                    return page_idx, ""

            # Use sequential processing for stability
            # PaddleOCR is not thread-safe and causes access violations (0xC0000005) when run in parallel
            logger.info("Starting sequential OCR processing (parallel disabled for stability)")
            
            for i, image in enumerate(images):
                pg_idx, pg_text = process_page(i, image)
                if pg_text.strip():
                    extracted_text_dict[pg_idx] = pg_text

            # Assemble text in order
            extracted_text = ""
            for i in range(pages_to_process):
                if i in extracted_text_dict and extracted_text_dict[i]:
                    extracted_text += f"\n--- Page {i+1} ---\n"
                    extracted_text += extracted_text_dict[i].strip() + "\n"

            logger.info(f"OCR completed. Extracted {len(extracted_text)} characters from {pages_to_process} pages")
            return extracted_text.strip()

        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            logger.error(traceback.format_exc())
            return ""

    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF using PyPDF2, fallback to OCR for image-based PDFs"""
        try:
            # First, try standard text extraction
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
            
            # Check if we got meaningful text
            cleaned_text = text.strip()
            word_count = len(cleaned_text.split()) if cleaned_text else 0
            
            # If we didn't get much text, it's likely an image-based PDF
            if word_count < 10 or len(cleaned_text) < 50:
                logger.info(f"Standard extraction yielded minimal text ({word_count} words). Trying fast OCR...")
                ocr_text = self.extract_text_with_ocr(file_path)

                if ocr_text and len(ocr_text.strip()) > len(cleaned_text):
                    logger.info("Fast OCR extraction successful, using OCR results")
                    return ocr_text
                else:
                    logger.info("OCR did not improve results, using standard extraction")
                    return cleaned_text
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            # If standard extraction fails completely, try PaddleOCR/Tesseract as last resort
            if PADDLEOCR_AVAILABLE or TESSERACT_AVAILABLE:
                logger.info("Standard extraction failed, attempting OCR as fallback")
                return self.extract_text_with_ocr(file_path)
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""
    
    def extract_text_from_doc(self, file_path):
        """Extract text from DOC files (basic XML parsing)"""
        try:
            # This is a very basic approach for .doc files
            # For better extraction, you'd need python-docx or antiword
            with open(file_path, 'rb') as file:
                content = file.read()
                # Try to find readable text (very basic approach)
                text = ""
                for byte in content:
                    if 32 <= byte <= 126:  # Printable ASCII characters
                        text += chr(byte)
                    elif byte in [10, 13]:  # Newlines
                        text += "\n"
                
                # Clean up the text
                lines = text.split('\n')
                cleaned_lines = []
                for line in lines:
                    line = line.strip()
                    if len(line) > 5 and not line.startswith('\\x'):
                        cleaned_lines.append(line)
                
                return '\n'.join(cleaned_lines)
        except Exception as e:
            logger.error(f"DOC extraction failed: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, file_path):
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            logger.error(f"Could not decode text file: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            return ""
    
    def clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters except newlines and tabs
        text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()

# Initialize text extractor
extractor = TextExtractor()

@app.route('/health', methods=['GET'])
def health_check():
    """Health check for text extraction service"""
    return jsonify({
        'status': 'healthy',
        'service': 'text_extraction',
        'supported_formats': ['pdf', 'docx', 'doc', 'txt'],
        'timestamp': datetime.now().isoformat()
    })

@app.route('/extract', methods=['POST'])
def extract_text():
    """Extract text from uploaded document"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get file type from filename or content type
        filename = file.filename.lower()
        content_type = file.content_type
        
        # Save temporary file
        temp_path = f"/tmp/{filename}"
        file.save(temp_path)
        
        try:
            text = ""
            
            # Determine file type and extract accordingly
            if filename.endswith('.pdf') or content_type == 'application/pdf':
                text = extractor.extract_text_from_pdf(temp_path)
            elif filename.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                text = extractor.extract_text_from_docx(temp_path)
            elif filename.endswith('.doc') or content_type == 'application/msword':
                text = extractor.extract_text_from_doc(temp_path)
            elif filename.endswith('.txt') or content_type == 'text/plain':
                text = extractor.extract_text_from_txt(temp_path)
            else:
                # Try as plain text
                text = extractor.extract_text_from_txt(temp_path)
            
            # Clean the extracted text
            cleaned_text = extractor.clean_text(text)
            
            # Get text statistics
            word_count = len(cleaned_text.split()) if cleaned_text else 0
            char_count = len(cleaned_text) if cleaned_text else 0
            
            return jsonify({
                'text': cleaned_text,
                'filename': filename,
                'content_type': content_type,
                'word_count': word_count,
                'character_count': char_count,
                'extraction_success': bool(cleaned_text),
                'timestamp': datetime.now().isoformat()
            })
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        logger.error(f"Text extraction error: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'error': 'Text extraction failed',
            'details': str(e)
        }), 500

@app.route('/extract/path', methods=['POST'])
def extract_text_from_path():
    """Extract text from file path (for internal Laravel use)"""
    try:
        data = request.get_json()
        
        if not data or 'file_path' not in data:
            return jsonify({'error': 'file_path is required'}), 400
        
        file_path = data['file_path']
        mime_type = data.get('mime_type', '')
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        text = ""
        filename = os.path.basename(file_path).lower()
        
        # Extract based on file extension or mime type
        if filename.endswith('.pdf') or mime_type == 'application/pdf':
            text = extractor.extract_text_from_pdf(file_path)
        elif filename.endswith('.docx') or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text = extractor.extract_text_from_docx(file_path)
        elif filename.endswith('.doc') or mime_type == 'application/msword':
            text = extractor.extract_text_from_doc(file_path)
        elif filename.endswith('.txt') or mime_type == 'text/plain':
            text = extractor.extract_text_from_txt(file_path)
        else:
            # Try as plain text
            text = extractor.extract_text_from_txt(file_path)
        
        # Clean the extracted text
        cleaned_text = extractor.clean_text(text)
        
        # Get text statistics
        word_count = len(cleaned_text.split()) if cleaned_text else 0
        char_count = len(cleaned_text) if cleaned_text else 0
        
        return jsonify({
            'text': cleaned_text,
            'file_path': file_path,
            'mime_type': mime_type,
            'word_count': word_count,
            'character_count': char_count,
            'extraction_success': bool(cleaned_text),
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Path-based text extraction error: {str(e)}")
        return jsonify({
            'error': 'Text extraction failed',
            'details': str(e)
        }), 500

if __name__ == '__main__':
    logger.info("Starting Text Extraction Service...")
    logger.info("Service will be available at: http://localhost:5002")
    logger.info("Health check: GET http://localhost:5002/health")
    logger.info("Extract from upload: POST http://localhost:5002/extract")
    logger.info("Extract from path: POST http://localhost:5002/extract/path")
    
    app.run(host='0.0.0.0', port=5002, debug=False, threaded=True)